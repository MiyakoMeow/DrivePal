"""Markdown 论文 → docx 转换器。一次性工具，用于学校提交格式转换。

用法: uv run python scripts/md2docx.py [-i input.md] [-o output.docx]
"""

from __future__ import annotations

import argparse
import asyncio
import hashlib
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

# ── 格式化常量（顶部集中可调） ─────────────────────────────────

# 字体
FONT_BODY = "宋体"
FONT_HEADING = "黑体"
FONT_CODE = "Courier New"

# 字号（磅）
SIZE_H1 = Pt(18)  # 二号
SIZE_H2 = Pt(16)  # 三号
SIZE_H3 = Pt(15)  # 小三
SIZE_H4 = Pt(14)  # 四号
SIZE_BODY = Pt(12)  # 小四
SIZE_CODE = Pt(10)  # 五号
SIZE_REF = Pt(9)  # 小五（参考文献）
SIZE_TOC = Pt(12)  # 目录正文

# 段落
LINE_SPACING = 1.5
FIRST_LINE_INDENT = Cm(0.85)  # 12pt 宋体约 2 字符

# 页面边距
MARGIN_TOP = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_LEFT = Cm(3.17)
MARGIN_RIGHT = Cm(3.17)

# mermaid
MERMAID_API = "https://mermaid.ink/img/"
MERMAID_TIMEOUT = 60  # 秒
MERMAID_CACHE_DIR = Path("archive/mermaid")

# LaTeX 公式渲染（codecogs API）
LATEX_API = "https://latex.codecogs.com/png.latex?"
LATEX_CACHE_DIR = Path("archive/latex")

# 默认路径
DEFAULT_INPUT = Path("archive/定稿-20260511.md")
DEFAULT_OUTPUT = Path("archive/定稿-20260511.docx")

# 引用正则（匹配文中引用标记 [N] 或 [N,M]）
_CITATION_RE = re.compile(r"\[\d+(?:,\d+)*\]")


# ── Markdown 解析 ─────────────────────────────────


async def parse(filepath: Path) -> list[dict]:
    """解析 markdown 文件为 token 字典列表。"""
    from markdown_it import MarkdownIt

    md = MarkdownIt()
    try:
        content = await asyncio.to_thread(filepath.read_text, encoding="utf-8")
        tokens = md.parse(content)
    except Exception as e:  # 宽捕获：任何解析失败均终止并退出，无需细分
        sys.exit(f"markdown 解析失败: {e}")

    result: list[dict] = []
    for t in tokens:
        d: dict = {"type": t.type, "tag": t.tag, "content": t.content}
        if t.attrs:
            d["attrs"] = dict(t.attrs)
        if hasattr(t, "info") and t.info:
            d["info"] = t.info
        if t.children:
            d["children"] = [{"type": c.type, "content": c.content} for c in t.children]
        result.append(d)
    return result


# ── 图片缓存校验 ─────────────────────────────────


def _mermaid_hash(code: str) -> str:
    """用 mermaid 代码内容的 SHA256 前 16 位作缓存文件名。"""
    return hashlib.sha256(code.encode()).hexdigest()[:16]


def _is_valid_png(path: Path) -> bool:
    """校验文件是否为有效 PNG。损坏则删除并返回 False。"""
    try:
        if path.stat().st_size < 8:
            path.unlink(missing_ok=True)
            return False
        with path.open("rb") as f:
            if f.read(8) != b"\x89PNG\r\n\x1a\n":
                path.unlink(missing_ok=True)
                return False
    except OSError:
        return False
    return True


# ── Mermaid / LaTeX 渲染 ─────────────────────────────────


async def _mkdir_cache(cache_dir: Path) -> None:
    """异步创建缓存目录。"""
    await asyncio.to_thread(cache_dir.mkdir, parents=True, exist_ok=True)


async def _write_cache(cache_path: Path, data: bytes) -> None:
    """异步写入缓存文件。"""
    await asyncio.to_thread(cache_path.write_bytes, data)


async def _check_cache(cache_path: Path) -> bool:
    """异步校验缓存 PNG 是否有效。"""
    return await asyncio.to_thread(_is_valid_png, cache_path)


async def _http_fetch_with_retry(client, url: str):
    """HTTP GET 带一次重试，仅对瞬态网络错误重试。"""
    import httpx

    for attempt in range(2):
        try:
            response = await client.get(url, timeout=MERMAID_TIMEOUT)
            response.raise_for_status()
            if response.content[:8] != b"\x89PNG\r\n\x1a\n":
                msg = "响应非 PNG 格式"
                raise ValueError(msg)
            return response.content
        except (
            httpx.TimeoutException,
            httpx.ConnectError,
            httpx.RemoteProtocolError,
            httpx.NetworkError,
        ):
            if attempt == 0:
                continue  # 一次重试
            raise
        except httpx.HTTPStatusError, ValueError:
            raise  # 非瞬态错误，不重试


async def render_mermaid(code: str, cache_dir: Path, client) -> Path | None:
    """将 mermaid 代码渲染为 PNG，缓存至 cache_dir。"""
    import json
    import zlib
    from base64 import urlsafe_b64encode

    payload = json.dumps({"code": code, "mermaid": '{"theme":"default"}'})
    data = zlib.compress(payload.encode(), level=9)
    encoded = urlsafe_b64encode(data).decode().rstrip("=")

    await _mkdir_cache(cache_dir)
    cache_path = cache_dir / f"{_mermaid_hash(code)}.png"
    if await _check_cache(cache_path):
        return cache_path

    url = f"{MERMAID_API}pako:{encoded}?type=png"
    try:
        data = await _http_fetch_with_retry(client, url)
        await _write_cache(cache_path, data)
        print(f"  [mermaid] 渲染成功 → {cache_path.name}")
        return cache_path
    except Exception as e:
        print(f"  [mermaid] 渲染失败: {e}", file=sys.stderr)
        return None


async def render_latex(
    latex: str, cache_dir: Path, client, *, display: bool = False
) -> Path | None:
    """将 LaTeX 公式渲染为 PNG，缓存至 cache_dir。"""
    import urllib.parse

    encoded = urllib.parse.quote(latex)
    extra = r"\dpi{200}" if display else r"\dpi{150}"
    url = f"{LATEX_API}{extra}{encoded}"

    await _mkdir_cache(cache_dir)
    cache_path = cache_dir / f"{_mermaid_hash(latex + str(display))}.png"
    if await _check_cache(cache_path):
        return cache_path

    try:
        data = await _http_fetch_with_retry(client, url)
        await _write_cache(cache_path, data)
        return cache_path
    except Exception as e:
        print(f"  [latex] 渲染失败: {e}", file=sys.stderr)
        return None


async def preprocess_tokens(tokens: list[dict], cache_dir: Path, client) -> list[dict]:
    """扫描 token 流，将 mermaid fence 替换为 image token。"""
    result: list[dict] = []
    i = 0
    while i < len(tokens):
        t = tokens[i]
        if t["type"] == "fence" and t.get("info", "").strip() == "mermaid":
            png_path = await render_mermaid(t["content"], cache_dir, client)
            if png_path:
                alt = ""
                peek = i + 1
                if peek < len(tokens) and tokens[peek]["type"] == "paragraph_open":
                    peek += 1
                if peek < len(tokens) and tokens[peek]["type"] == "inline":
                    next_content = tokens[peek].get("content", "")
                    if next_content.strip().startswith("**图"):
                        alt = next_content.strip().strip("*")
                result.append({"type": "image", "path": str(png_path), "alt": alt})
            else:
                result.append(
                    {
                        "type": "paragraph",
                        "content": f"[图表渲染失败，请手动处理]\n```mermaid\n{t['content']}\n```",
                    }
                )
            i += 1
            peek = i
            if (
                peek < len(tokens)
                and tokens[peek]["type"] == "paragraph_open"
                and peek + 1 < len(tokens)
                and tokens[peek + 1]["type"] == "inline"
                and peek + 2 < len(tokens)
                and tokens[peek + 2]["type"] == "paragraph_close"
                and tokens[peek + 1].get("content", "").strip().startswith("**图")
            ):
                i += 3
            continue
        result.append(t)
        i += 1
    return result


# ── docx 构建 ─────────────────────────────────


async def _check_font_warning() -> None:
    """检测中文字体可用性，缺失时警告。"""
    import platform

    def _scan() -> None:
        system = platform.system()
        if system == "Linux":
            font_dirs = ["/usr/share/fonts", "/usr/local/share/fonts"]
        elif system == "Darwin":
            font_dirs = ["/System/Library/Fonts", "/Library/Fonts"]
        else:
            return

        for font_name, label in [(FONT_BODY, "正文"), (FONT_HEADING, "标题")]:
            found = False
            for fd in font_dirs:
                if not os.path.isdir(fd):
                    continue
                for _root, _dirs, files in os.walk(fd):
                    if any(font_name.lower() in f.lower() for f in files):
                        found = True
                        break
                if found:
                    break
            if not found:
                print(
                    f"警告: 字体 '{font_name}'（{label}）未在系统中找到，将使用 Calibri 回退",
                    file=sys.stderr,
                )

    await asyncio.to_thread(_scan)


def _setup_document(doc: Document) -> None:
    """配置页面边距、默认字体。字体警告由调用方异步执行。"""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT

    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = SIZE_BODY
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_BODY)


def _add_run(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    font_name: str | None = None,
    font_size: Pt | None = None,
    superscript: bool = False,
) -> None:
    """添加一个格式化的 run 到段落。"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if superscript:
        run.font.superscript = True
    if font_name:
        run.font.name = font_name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)
    if font_size:
        run.font.size = font_size


def _new_paragraph(
    doc: Document,
    text: str = "",
    *,
    alignment: int | None = None,
    font_name: str | None = None,
    font_size: Pt | None = None,
    bold: bool = False,
    spacing: float | None = None,
    first_line_indent: Cm | None = None,
):
    """创建新段落，可选设置对齐/字体/行距/首行缩进。"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    if spacing is not None:
        pf.line_spacing = spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if text:
        _add_run(p, text, bold=bold, font_name=font_name, font_size=font_size)
    return p


# ── 标题处理 ──


def _add_heading(doc: Document, text: str, level: int) -> None:
    """添加标题段落。"""
    sizes = {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3, 4: SIZE_H4}
    alignments = {1: WD_ALIGN_PARAGRAPH.CENTER}
    spacing_after = {1: Pt(24), 2: Pt(18), 3: Pt(12), 4: Pt(8)}

    p = _new_paragraph(
        doc,
        text,
        bold=True,
        font_name=FONT_HEADING,
        font_size=sizes.get(level, SIZE_H4),
        alignment=alignments.get(level, WD_ALIGN_PARAGRAPH.LEFT),
        spacing=LINE_SPACING,
    )
    p.paragraph_format.space_after = spacing_after.get(level, Pt(6))
    # 学校提交格式要求每章另起一页
    if level == 2:
        p.paragraph_format.page_break_before = True


def _extract_plain_text(md_content: str) -> str:
    """从 markdown inline 文本中提取纯文本（去除粗斜体标记）。"""
    from markdown_it import MarkdownIt

    md = MarkdownIt()
    inline_tokens = md.parseInline(md_content)
    return "".join(t.content for t in inline_tokens if t.content)


# ── 正文处理 ──


async def _render_inline_content(
    paragraph, content: str, *, client, font_size: Pt = SIZE_BODY
) -> None:
    """将 inline markdown 文本渲染为 paragraph 中的 runs。

    支持：**bold**/__bold__、*italic*/_italic_、`code`、[N] 上标引用、$...$ 行内公式。
    """
    i = 0
    n = len(content)
    while i < n:
        # 上标引用 [N] 或 [N,M]
        if content[i] == "[":
            m = _CITATION_RE.match(content, i)
            if m:
                _add_run(paragraph, m.group(), superscript=True, font_size=font_size)
                i = m.end()
                continue
        # 粗体 **...** 或 __...__
        if i + 1 < n and content[i : i + 2] in ("**", "__"):
            marker = content[i : i + 2]
            end = content.find(marker, i + 2)
            if end != -1:
                _render_inner_with_citations(
                    paragraph, content[i + 2 : end], bold=True, font_size=font_size
                )
                i = end + 2
                continue
            # 无闭合标记 → 按纯文本输出，避免静默丢弃
            _add_run(paragraph, marker, font_size=font_size)
            i += 2
            continue
        # 斜体 *...* 或 _..._（需排除 __ 情况）
        if content[i] in "*_":
            if content[i] == "_" and i + 1 < n and content[i + 1] == "_":
                _add_run(paragraph, content[i], font_size=font_size)
                i += 1
                continue
            end = content.find(content[i], i + 1)
            if end != -1:
                _render_inner_with_citations(
                    paragraph, content[i + 1 : end], italic=True, font_size=font_size
                )
                i = end + 1
                continue
        # 行内代码 `...`
        if content[i] == "`":
            end = content.find("`", i + 1)
            if end != -1:
                _add_run(
                    paragraph,
                    content[i + 1 : end],
                    font_name=FONT_CODE,
                    font_size=SIZE_CODE,
                )
                i = end + 1
                continue
        # 行内公式 $...$（不匹配 $$ 块级公式）
        if content[i] == "$" and (i + 1 >= n or content[i + 1] != "$"):
            end = content.find("$", i + 1)
            if end != -1:
                latex = content[i + 1 : end]
                png_path = await render_latex(latex, LATEX_CACHE_DIR, client)
                if png_path:
                    run = paragraph.add_run()
                    run.add_picture(str(png_path), height=Inches(0.22))
                else:
                    _add_run(paragraph, f"[公式:{latex}]", font_size=SIZE_CODE)
                i = end + 1
                continue
        # 普通文本
        j = i
        while j < n and content[j] not in "*[_`$":
            j += 1
        if j > i:
            _add_run(paragraph, content[i:j], font_size=font_size)
            i = j
            continue
        # 未匹配的特殊字符
        _add_run(paragraph, content[i], font_size=font_size)
        i += 1


def _render_inner_with_citations(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    font_size: Pt = SIZE_BODY,
) -> None:
    """渲染粗/斜体内部文本，同时解析其中的 [N] 引用上标。"""
    parts = _CITATION_RE.split(text)
    citations = _CITATION_RE.findall(text)
    for k, part in enumerate(parts):
        if part:
            _add_run(paragraph, part, bold=bold, italic=italic, font_size=font_size)
        if k < len(citations):
            _add_run(
                paragraph,
                citations[k],
                bold=bold,
                italic=italic,
                superscript=True,
                font_size=font_size,
            )


async def _add_body_paragraph(
    doc: Document, content: str, *, client, use_ref_font: bool = False
) -> None:
    """添加正文段落，处理粗斜体、[N] 引用和公式。"""
    # 块级公式 $$...$$：整段替换为居中图片
    stripped = content.strip()
    if stripped.startswith("$$") and stripped.endswith("$$"):
        latex = stripped[2:-2].strip()
        png_path = await render_latex(latex, LATEX_CACHE_DIR, client, display=True)
        if png_path:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(png_path), width=Inches(5.5))
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p, f"[公式渲染失败: {latex}]", font_size=SIZE_CODE)
        return

    font_size = SIZE_REF if use_ref_font else SIZE_BODY
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    await _render_inline_content(p, content, client=client, font_size=font_size)


# ── 代码块 ──


def _add_code_block(doc: Document, code: str) -> None:
    """添加等宽字体代码块段落。"""
    for line in code.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Cm(0)
        _add_run(p, line, font_name=FONT_CODE, font_size=SIZE_CODE)


# ── 表格 ──


def _render_table(doc: Document, rows: list[list[str]]) -> None:
    """将二维数据渲染为三线表。首行（表头）加粗。"""
    if not rows or not rows[0]:
        return

    expected_cols = len(rows[0])
    for row_data in rows:
        while len(row_data) < expected_cols:
            row_data.append("")

    table = doc.add_table(rows=len(rows), cols=expected_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        is_header = i == 0
        for j, cell_text in enumerate(row_data):
            if j >= len(row.cells):
                break
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            # 剥离表格单元格内 Markdown 标记符（粗斜体/代码/引用等）
            plain = _extract_plain_text(cell_text)
            _add_run(p, plain, font_size=SIZE_CODE, bold=is_header)
            p.paragraph_format.line_spacing = 1.0

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = tbl.makeelement(qn("w:tblPr"), {})
        tbl.insert(0, tblPr)  # w:tblPr 须为 w:tbl 首个子元素
    # 移除样式可能预设的 tblBorders，避免重复元素
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    top = borders.makeelement(
        qn("w:top"), {qn("w:val"): "single", qn("w:sz"): "12", qn("w:color"): "000000"}
    )
    borders.append(top)
    bottom = borders.makeelement(
        qn("w:bottom"),
        {qn("w:val"): "single", qn("w:sz"): "12", qn("w:color"): "000000"},
    )
    borders.append(bottom)
    insideH = borders.makeelement(
        qn("w:insideH"),
        {qn("w:val"): "single", qn("w:sz"): "4", qn("w:color"): "000000"},
    )
    borders.append(insideH)
    insideV = borders.makeelement(
        qn("w:insideV"), {qn("w:val"): "none", qn("w:sz"): "0"}
    )
    borders.append(insideV)
    left = borders.makeelement(
        qn("w:left"), {qn("w:val"): "single", qn("w:sz"): "4", qn("w:color"): "000000"}
    )
    borders.append(left)
    right = borders.makeelement(
        qn("w:right"), {qn("w:val"): "single", qn("w:sz"): "4", qn("w:color"): "000000"}
    )
    borders.append(right)
    tblPr.append(borders)


# ── 列表 ──


async def _add_list_item(
    doc: Document, content: str, client, is_ordered: bool, counter: int
) -> None:
    """添加列表项段落。"""
    prefix = f"{counter}. " if is_ordered else "\u2022 "
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.85)
    await _render_inline_content(
        p, prefix + content, client=client, font_size=SIZE_BODY
    )


# ── 目录和页码 ──


def _collect_headings(tokens: list[dict]) -> list[tuple[int, str]]:
    """收集所有 h2/h3 标题（层级, 文本）。"""
    headings: list[tuple[int, str]] = []
    pending_level: int | None = None
    for t in tokens:
        if t["type"] == "heading_open":
            pending_level = int(t["tag"][1])
            continue
        if t["type"] == "heading_close":
            pending_level = None
            continue
        if t["type"] == "inline" and pending_level is not None and pending_level >= 2:
            text = _extract_plain_text(t["content"])
            headings.append((pending_level, text))
            pending_level = None
    return headings


def _add_toc(doc: Document, headings: list[tuple[int, str]]) -> None:
    """插入静态目录页。

    仅含标题与缩进，不含页码。python-docx 不支持 TOC 域代码，
    静态目录为已知限制——转换后在 Word 中手动补页码或插入自动目录。
    """
    p = _new_paragraph(
        doc,
        "目录",
        bold=True,
        font_name=FONT_HEADING,
        font_size=SIZE_H2,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    p.paragraph_format.space_after = Pt(18)

    for level, text in headings:
        indent = Cm(0) if level == 2 else Cm(0.85)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = indent
        _add_run(p, text, font_size=SIZE_TOC)

    doc.add_page_break()


def _add_page_numbers(doc: Document) -> None:
    """在页脚居中添加页码。"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = run._element.makeelement(
            qn("w:fldChar"), {qn("w:fldCharType"): "begin"}
        )
        run._element.append(fldChar1)
        run2 = p.add_run()
        instrText = run2._element.makeelement(
            qn("w:instrText"), {qn("xml:space"): "preserve"}
        )
        instrText.text = " PAGE "
        run2._element.append(instrText)
        run3 = p.add_run()
        fldChar2 = run3._element.makeelement(
            qn("w:fldChar"), {qn("w:fldCharType"): "end"}
        )
        run3._element.append(fldChar2)


# ── 主构建函数 ──


async def build_docx(tokens: list[dict], output_path: Path, client) -> None:
    """主建文档函数。遍历 token 流构建 docx。"""
    doc = Document()
    _setup_document(doc)

    # 目录
    headings = _collect_headings(tokens)
    _add_toc(doc, headings)

    # 状态变量
    in_references = False
    pending_heading_level: int | None = None
    in_table = False
    table_rows: list[list[str]] = []
    current_row: list[str] = []
    in_cell = False
    current_cell: list[str] = []
    in_list = False
    ordered_list = False
    list_counter = 0
    pending_list_item = False

    for idx, t in enumerate(tokens):
        tp = t["type"]

        # ── 标题 ──
        if tp == "heading_open":
            pending_heading_level = int(t["tag"][1])
            continue
        if tp == "heading_close":
            pending_heading_level = None
            continue
        if tp == "inline" and pending_heading_level is not None:
            text = _extract_plain_text(t["content"])
            _add_heading(doc, text, pending_heading_level)
            if "参考文献" in text:
                in_references = True
            pending_heading_level = None
            continue

        # ── 水平分隔线 → 分页 ──
        if tp == "hr":
            doc.add_page_break()
            continue

        # ── 图片（mermaid 预处理后） ──
        if tp == "image":
            png_path = Path(t["path"])
            if await asyncio.to_thread(png_path.exists):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(png_path), width=Inches(6))
            else:
                await _add_body_paragraph(
                    doc, f"[图片缺失: {png_path.name}]", client=client
                )
            continue

        # ── 代码块（非 mermaid） ──
        if tp == "fence" and t.get("info", "").strip() != "mermaid":
            _add_code_block(doc, t["content"])
            continue

        # ── 表格 ──
        if tp == "table_open":
            in_table = True
            table_rows = []
            in_cell = False
            current_cell = []
            continue
        if tp == "table_close":
            in_table = False
            _render_table(doc, table_rows)
            table_rows = []
            continue
        if tp == "tr_open":
            current_row = []
            continue
        if tp == "tr_close":
            if current_row:
                table_rows.append(current_row)
            current_row = []
            continue
        if tp in ("th_open", "td_open"):
            in_cell = True
            current_cell = []
            continue
        if tp in ("th_close", "td_close"):
            in_cell = False
            current_row.append("".join(current_cell))
            current_cell = []
            continue
        if tp == "inline" and in_table and in_cell:
            current_cell.append(t["content"])
            continue

        # ── 列表 ──
        if tp == "bullet_list_open":
            in_list = True
            ordered_list = False
            list_counter = 0
            continue
        if tp == "bullet_list_close":
            in_list = False
            continue
        if tp == "ordered_list_open":
            in_list = True
            ordered_list = True
            list_counter = 0
            continue
        if tp == "ordered_list_close":
            in_list = False
            continue
        if tp == "list_item_open":
            pending_list_item = True
            list_counter += 1
            continue
        if tp == "list_item_close":
            pending_list_item = False
            continue
        if tp == "inline" and in_list and pending_list_item:
            await _add_list_item(doc, t["content"], client, ordered_list, list_counter)
            pending_list_item = False
            continue

        # ── 正文段落 ──
        if tp == "paragraph_open":
            continue
        if tp == "paragraph_close":
            continue
        if tp == "inline":
            await _add_body_paragraph(
                doc, t["content"], client=client, use_ref_font=in_references
            )
            continue
        if tp in ("hardbreak", "softbreak"):
            continue

    # 页码
    _add_page_numbers(doc)

    if await asyncio.to_thread(output_path.exists):
        print(f"覆盖已有文件: {output_path}", file=sys.stderr)

    def _save() -> None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

    await asyncio.to_thread(_save)
    print(f"输出: {output_path}")


# ── 入口 ──


async def main() -> None:
    import httpx

    parser = argparse.ArgumentParser(description="Markdown 论文 → docx 转换器")
    parser.add_argument(
        "-i", "--input", type=Path, default=DEFAULT_INPUT, help="输入 markdown 文件路径"
    )
    parser.add_argument(
        "-o", "--output", type=Path, default=DEFAULT_OUTPUT, help="输出 docx 文件路径"
    )
    args = parser.parse_args()

    if not args.input.exists():
        sys.exit(f"输入文件不存在: {args.input}")

    async with httpx.AsyncClient() as client:
        await _check_font_warning()

        print(f"解析: {args.input}")
        tokens = await parse(args.input)
        print(f"  获得 {len(tokens)} 个 token")

        print("渲染 Mermaid 图表...")
        tokens = await preprocess_tokens(tokens, MERMAID_CACHE_DIR, client)
        print(f"  预处理后 {len(tokens)} 个 token")

        print("构建 docx...")
        await build_docx(tokens, args.output, client)


if __name__ == "__main__":
    asyncio.run(main())
